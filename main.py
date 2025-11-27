import os
import re
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import win32com.client  # 用于格式转换
from docx import Document  # 用于docx文档基本操作
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 用于段落对齐设置
from docx.oxml import OxmlElement  # 用于操作XML元素
from docx.oxml.ns import qn  # 用于设置XML命名空间
from docx.shared import Pt, Cm  # 用于设置字体大小和厘米单位
import concurrent.futures  # 用于线程池并行处理
import threading  # 用于线程锁和线程管理

# ------------------------------
# 全局变量（并行处理+正则缓存）
# ------------------------------
# 正则缓存（全局编译，只解析一次，提升速度）
english_pattern = re.compile(r'\(([012]\d)[^()]*?[\u4e00-\u9fa5][^()]{0,10}\)')  # 英文小括号
chinese_pattern = re.compile(r'（([012]\d)[^（）]*?[\u4e00-\u9fa5][^（）]{0,10}）')  # 中文小括号
k_pattern = re.compile(r'\[([012]\d)[^\]]*?[\u4e00-\u9fa5][^\]]{0,10}\]')  # 英文中括号

# 并行处理进度统计（线程安全）
progress_lock = threading.Lock()
processed_count = 0  # 已处理文件数
success_count = 0  # 成功文件数
error_list = []  # 错误列表
options = {}  # 全局配置选项
root = None  # 主窗口对象
process_btn = None  # 处理按钮对象
status_var = None  # 状态显示变量
folder_var = None  # 文件夹路径变量
convert_doc_btn = None  # DOC转DOCX按钮
convert_pdf_btn = None  # DOCX转PDF按钮
total_files = 0  # 并行处理总文件数


# ------------------------------
# 通用工具函数
# ------------------------------
def select_folder():
    """选择文件夹并返回路径，若取消选择则返回None"""
    folder_path = filedialog.askdirectory(title="选择文件夹")
    return folder_path if folder_path else None


def get_all_files_by_ext(folder_path, exts):
    """
    获取指定文件夹下所有符合扩展名的文件路径（排除Word临时文件）
    :param folder_path: 文件夹路径
    :param exts: 扩展名列表（如['.docx', '.doc']）
    :return: 符合条件的文件路径列表
    """
    exts = [ext.lower() for ext in exts]  # 统一转为小写便于匹配
    files = []
    # 遍历文件夹及其子文件夹
    for root_dir, _, filenames in os.walk(folder_path):
        for filename in filenames:
            # 过滤Word临时文件（以~$开头的文件）
            if filename.startswith('~$'):
                continue
            # 检查文件是否符合任一扩展名
            if any(filename.lower().endswith(ext) for ext in exts):
                files.append(os.path.join(root_dir, filename))
    return files


# ------------------------------
# 主功能区：Word处理功能
# ------------------------------
def remove_header_footer(doc):
    """删除文档中所有节的页眉页脚内容，并断开节链接"""
    for section in doc.sections:
        # 关键：断开当前节与前一节的页眉页脚链接
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # 原有删除页眉逻辑
        header = section.header
        for para in reversed(header.paragraphs):
            p_element = para._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            para._p = None
            para._element = None

        # 原有删除页脚逻辑
        footer = section.footer
        for para in reversed(footer.paragraphs):
            p_element = para._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            para._p = None
            para._element = None


def add_custom_header(doc):
    """为文档所有节添加自定义页眉，距离顶端0.7cm，避免重复"""
    for section in doc.sections:
        # 1. 设置页眉距离顶端 0.7cm
        section.header_distance = Cm(0.7)

        # 2. 强制清空当前节页眉（去重逻辑）
        header = section.header
        for para in reversed(header.paragraphs):
            p_element = para._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            para._p = None
            para._element = None

        # 3. 添加新页眉内容
        para = header.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.left_indent = Cm(2) - section.left_margin
        run = para.add_run("泉尚优学：学为人师，行为世范！")
        run.font.name = "华文行楷"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "华文行楷")
        run.font.size = Pt(12)


def add_centered_page_number(doc):
    """为文档所有节添加居中页码（第X页/共Y页），距离底端1cm，避免重复"""
    for section in doc.sections:
        # 1. 设置页脚距离底端 1cm
        section.footer_distance = Cm(1.0)

        # 2. 强制清空当前节页脚（去重逻辑）
        footer = section.footer
        for para in reversed(footer.paragraphs):
            p_element = para._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            para._p = None
            para._element = None

        # 3. 添加页码内容
        p = footer.add_paragraph()
        para_format = p.paragraph_format
        para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("第")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        run.font.size = Pt(12)

        run = p.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"
        run._r.append(instr_text)

        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fld_char_sep)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)

        run = p.add_run("页/共")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        run.font.size = Pt(12)

        run = p.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.text = "NUMPAGES"
        run._r.append(instr_text)

        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fld_char_sep)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)

        run = p.add_run("页")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        run.font.size = Pt(12)


def replace_patterns_in_paragraph(paragraph):
    """
    替换段落中符合特定模式的文本（复用全局编译的正则，提升速度）
    :param paragraph: 需要处理的段落对象
    """
    text_runs = []  # 存储段落中所有文本片段（包含run对象、文本内容及位置）
    char_pos = 0  # 字符位置计数器

    for run in paragraph.runs:
        if not run.text:
            continue  # 跳过空文本
        text = run.text
        start = char_pos
        end = char_pos + len(text)
        text_runs.append((run, text, start, end))
        char_pos = end  # 更新位置

    if not text_runs:
        return  # 无文本则直接返回

    # 合并所有文本用于匹配
    all_text = ''.join([t[1] for t in text_runs])
    replaced_ranges = []  # 存储需要替换的文本范围

    # 标记需要替换的范围
    def mark_replaced(match):
        replaced_ranges.append((match.start(), match.end()))
        return ""  # 替换为空

    # 执行匹配并标记（复用全局正则，无需重复编译）
    re.sub(chinese_pattern, mark_replaced, all_text)
    re.sub(english_pattern, mark_replaced, all_text)
    re.sub(k_pattern, mark_replaced, all_text)

    # 创建保留标记（True表示保留，False表示删除）
    keep_mask = [True] * len(all_text)
    for start, end in replaced_ranges:
        for i in range(start, end):
            if i < len(keep_mask):
                keep_mask[i] = False

    # 根据保留标记更新每个run的文本
    for run, original_text, start, end in text_runs:
        kept_chars = []
        for i in range(start, end):
            if i < len(keep_mask) and keep_mask[i]:
                original_idx = i - start  # 计算在原始文本中的索引
                kept_chars.append(original_text[original_idx])
        run.text = ''.join(kept_chars)  # 更新run的文本


def set_outline_level(doc):
    """
    将文档中符合特定格式的段落大纲级别设置为1级
    """
    # 定义中文数字（扩展常用范围）
    chinese_nums = r'(?:一|二|三|四|五|六|七|八|九|十|十一|十二|十三|十四|十五|十六|十七|十八|十九|二十)'

    # 组合所有匹配模式
    pattern = (
            r'^\s*('
            r'(题型|考点|考法)(?:\d+|' + chinese_nums + r').*'
            r'|(?:A夯实基础|B能力提升|C综合素养)\s*$'
            r'|第(?:\d+|' + chinese_nums + r')(章|单元).*'
            r')'
    )

    for para in doc.paragraphs:
        clean_text = para.text.strip()
        if re.search(pattern, clean_text):
            # 获取或创建段落属性元素
            p_pr = para._element.get_or_add_pPr()
            # 移除已有的大纲级别设置（避免重复）
            for elem in p_pr.findall(qn('w:outlineLvl')):
                p_pr.remove(elem)
            # 创建大纲级别元素并设置为1级（Word中0对应1级）
            outline_level = OxmlElement('w:outlineLvl')
            outline_level.set(qn('w:val'), '0')
            p_pr.append(outline_level)


def process_word_file(file_path, keep_backup):
    """
    处理单个Word文件（根据选项执行相应操作）
    :param file_path: 文件路径
    :param keep_backup: 是否保留备份
    :return: (处理结果, 消息)
    """
    try:
        # 保留备份（若未存在备份）
        if keep_backup and not os.path.exists(f"{file_path}.bak"):
            shutil.copy2(file_path, f"{file_path}.bak")

        # 打开文档进行处理
        doc = Document(file_path)

        # 根据选项执行操作
        if options['remove_header_footer'].get():
            remove_header_footer(doc)
        if options['add_page_number'].get():
            add_centered_page_number(doc)
        if options['add_custom_header'].get():
            add_custom_header(doc)

        if options['replace_patterns'].get():
            # 处理普通段落
            for para in doc.paragraphs:
                replace_patterns_in_paragraph(para)
            # 处理表格中的段落
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_patterns_in_paragraph(para)

        # 设置题型段落的大纲级别为1级
        if options['set_question_outline'].get():
            set_outline_level(doc)

        # 保存修改
        doc.save(file_path)
        return True, f"成功：{os.path.basename(file_path)}"
    except Exception as e:
        return False, f"失败：{os.path.basename(file_path)} - {str(e)}"


# ------------------------------
# 并行处理核心逻辑（原有Word处理）
# ------------------------------
def process_single_file(file_path, keep_backup):
    """单个文件的处理逻辑（线程池执行单元，线程安全）"""
    global processed_count, success_count, error_list
    filename = os.path.basename(file_path)
    try:
        # 执行文件处理
        res, msg = process_word_file(file_path, keep_backup)
        # 线程安全更新统计数据（用锁避免并发冲突）
        with progress_lock:
            processed_count += 1
            if res:
                success_count += 1
            else:
                error_list.append(msg)
        # 实时更新UI进度（通过主线程after方法，确保UI安全）
        root.after(0, lambda: status_var.set(
            f"并行处理中 ({processed_count}/{total_files})：当前处理 {filename}"
        ))
    except Exception as e:
        # 捕获未知错误
        with progress_lock:
            processed_count += 1
            error_list.append(f"失败：{filename} - 未知错误：{str(e)}")
        root.after(0, lambda: status_var.set(
            f"并行处理中 ({processed_count}/{total_files})：{filename} 处理失败"
        ))


def finish_process(keep_backup):
    """所有文件处理完成后，显示结果并恢复UI"""
    global processed_count, success_count, error_list
    result = f"并行处理完成！\n成功：{success_count}/{processed_count}\n"
    if keep_backup:
        result += "原文件已备份为.bak格式，处理后的文件已替换原文件"
    else:
        result += "已直接替换原文件（未保留备份）"

    if error_list:
        result += f"\n\n错误列表（前5条）：\n" + "\n".join(error_list[:5])
    messagebox.showinfo("并行处理结果", result)
    # 恢复按钮和状态
    process_btn.config(state=tk.NORMAL)
    status_var.set("就绪")
    # 重置全局进度统计
    with progress_lock:
        processed_count = 0
        success_count = 0
        error_list = []


def start_parallel_process():
    """启动线程池并行处理（子线程中执行，不阻塞UI）"""
    global total_files
    folder_path = folder_var.get().replace("已选择：", "")
    keep_backup = options['keep_backup'].get()
    word_files = get_all_files_by_ext(folder_path, ['.docx'])
    total_files = len(word_files)

    # 配置线程池大小：IO密集型任务最优为 CPU核心数*2，最多10个线程避免资源占用过高
    max_workers = min(10, total_files)
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 批量提交任务（每个文件一个任务）
        executor.map(
            process_single_file,  # 任务执行函数
            word_files,  # 第一个参数：文件路径列表
            [keep_backup] * total_files  # 第二个参数：是否保留备份（每个任务相同）
        )

    # 所有任务完成后，调用收尾函数
    root.after(0, lambda: finish_process(keep_backup))


def process_word_files_action():
    """处理Word文件的入口函数（启动子线程，避免阻塞UI）"""
    global processed_count, success_count, error_list
    # 重置全局进度统计
    with progress_lock:
        processed_count = 0
        success_count = 0
        error_list = []

    folder_path = folder_var.get().replace("已选择：", "")
    if not folder_path or folder_path == "等待选择文件夹...":
        messagebox.showwarning("警告", "请先选择文件夹")
        return

    # 检查是否选择了至少一个处理选项
    if not any([
        options['remove_header_footer'].get(),
        options['add_custom_header'].get(),
        options['add_page_number'].get(),
        options['replace_patterns'].get(),
        options['set_question_outline'].get()
    ]):
        if not messagebox.askyesno("提示", "未选择任何处理选项，是否继续？"):
            return

    word_files = get_all_files_by_ext(folder_path, ['.docx'])
    total = len(word_files)
    if total == 0:
        messagebox.showinfo("提示", "未找到任何.docx文件")
        return

    # 禁用按钮+更新状态
    process_btn.config(state=tk.DISABLED)
    status_var.set(f"准备并行处理 {total} 个文件...")
    root.update_idletasks()

    # 启动子线程执行并行处理（守护线程，避免程序退出残留）
    thread = threading.Thread(target=start_parallel_process, daemon=True)
    thread.start()


# ------------------------------
# 辅助功能区：格式转换功能（并行优化版）
# ------------------------------
def split_tasks(file_list, thread_count):
    """均分任务到线程（新增函数）"""
    batch_size = (len(file_list) + thread_count - 1) // thread_count  # 向上取整
    return [file_list[i * batch_size: min((i + 1) * batch_size, len(file_list))] for i in range(thread_count)]


def show_convert_result(convert_type, total, extra_params):
    """显示转换结果（新增函数，适配并行统计）"""
    global processed_count, success_count, error_list
    result_msg = f"{convert_type} 并行处理完成！\n总文件：{total}\n成功转换：{success_count}\n"

    # 统计跳过数（排除真正的错误）
    skipped_count = total - success_count - len([e for e in error_list if "跳过" not in e])
    result_msg += f"跳过（已存在/无需处理）：{skipped_count}\n"

    # 根据转换类型补充信息
    if convert_type == "DOC→DOCX":
        keep_source = extra_params
        result_msg += f"保留源文件：{'是' if keep_source else '否'}\n"
    elif convert_type == "DOCX→PDF":
        use_separate_folder = extra_params
        save_path = "独立的 docx2pdf 文件夹" if use_separate_folder else "原文件所在位置"
        result_msg += f"PDF保存位置：{save_path}\n"

    # 追加错误信息
    if error_list:
        result_msg += "\n错误详情（前5条）：\n" + "\n".join(error_list[:5])
    root.after(0, lambda: messagebox.showinfo(f"{convert_type} 结果", result_msg))

    # 恢复UI状态
    root.after(0, lambda: status_var.set("就绪"))
    if convert_type == "DOC→DOCX":
        root.after(0, lambda: convert_doc_btn.config(state=tk.NORMAL))
    else:
        root.after(0, lambda: convert_pdf_btn.config(state=tk.NORMAL))


def parallel_convert_doc_to_docx(root_dir, keep_source, status_var):
    """并行批量将doc文件转换为docx文件（替换原 batch_convert_doc_to_docx 函数）"""
    global processed_count, success_count, error_list
    # 重置统计变量（线程安全）
    with progress_lock:
        processed_count = 0
        success_count = 0
        error_list = []

    root_dir = os.path.normpath(root_dir)
    # 1. 收集所有待转换的doc文件（排除docx、临时文件）
    doc_files = []
    for foldername, _, filenames in os.walk(root_dir):
        for filename in filenames:
            lower_name = filename.lower()
            if lower_name.endswith(".doc") and not lower_name.endswith(".docx") and not filename.startswith("~$"):
                doc_files.append(os.path.join(foldername, filename))
    total = len(doc_files)
    if total == 0:
        root.after(0,
                   lambda: [messagebox.showinfo("提示", "未找到任何.doc文件"), convert_doc_btn.config(state=tk.NORMAL)])
        return

    # 2. 配置线程数（最多5个，避免Word进程过多）
    cpu_count = os.cpu_count() or 2
    max_threads = min(5, cpu_count * 1, total)
    task_queues = split_tasks(doc_files, max_threads)  # 均分任务

    # 3. 单个线程的工作逻辑（每个线程独占一个Word实例）
    def worker(task_queue):
        nonlocal keep_source
        # 线程内创建独立的Word进程
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        try:
            for doc_path in task_queue:
                filename = os.path.basename(doc_path)
                # 异步更新UI进度
                root.after(0, lambda f=filename: status_var.set(
                    f"并行转换DOC→DOCX（{processed_count + 1}/{total}）：{f}"
                ))

                # 构建目标docx路径
                docx_name = f"{os.path.splitext(filename)[0]}.docx"
                docx_path = os.path.join(os.path.dirname(doc_path), docx_name)

                success_flag = False
                error_msg = ""
                # 目标文件不存在才转换
                if not os.path.exists(docx_path):
                    try:
                        doc = word.Documents.Open(os.path.abspath(doc_path))
                        doc.SaveAs2(os.path.abspath(docx_path), FileFormat=12)  # 12=docx格式
                        doc.Close()
                        success_flag = True
                    except Exception as e:
                        error_msg = f"转换失败：{str(e).split(',')[0]}"
                else:
                    error_msg = "跳过：目标docx已存在"

                # 线程安全更新统计数据
                with progress_lock:
                    global processed_count, success_count, error_list
                    processed_count += 1
                    if success_flag:
                        success_count += 1
                    if error_msg:
                        error_list.append(f"{filename} - {error_msg}")

                # 不需要保留源文件则删除
                if success_flag and not keep_source:
                    try:
                        if os.path.exists(doc_path):
                            os.remove(doc_path)
                    except Exception as e:
                        with progress_lock:
                            error_list.append(f"{filename} - 删除源文件失败：{str(e)}")
        finally:
            # 必须关闭Word进程，释放资源
            word.Quit()

    # 4. 启动线程池执行任务
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_threads) as executor:
        executor.map(worker, task_queues)

    # 5. 任务完成后显示结果
    root.after(0, lambda: show_convert_result("DOC→DOCX", total, keep_source))


def parallel_convert_docx_to_pdf(root_dir, use_separate_folder, status_var):
    """并行批量将docx文件转换为pdf文件（替换原 batch_convert_docx_to_pdf 函数）"""
    global processed_count, success_count, error_list
    # 重置统计变量（线程安全）
    with progress_lock:
        processed_count = 0
        success_count = 0
        error_list = []

    root_dir = os.path.normpath(root_dir)
    # 1. 收集所有待转换的docx文件（排除临时文件）
    docx_files = get_all_files_by_ext(root_dir, ['.docx'])
    total = len(docx_files)
    if total == 0:
        root.after(0, lambda: [messagebox.showinfo("提示", "未找到任何.docx文件"),
                               convert_pdf_btn.config(state=tk.NORMAL)])
        return

    # 2. 配置线程数（最多5个，避免Word进程过多）
    cpu_count = os.cpu_count() or 2
    max_threads = min(5, cpu_count * 1, total)
    task_queues = split_tasks(docx_files, max_threads)  # 均分任务

    # 3. 单个线程的工作逻辑（每个线程独占一个Word实例）
    def worker(task_queue):
        nonlocal use_separate_folder
        # 线程内创建独立的Word进程
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        try:
            for docx_path in task_queue:
                filename = os.path.basename(docx_path)
                # 异步更新UI进度
                root.after(0, lambda f=filename: status_var.set(
                    f"并行转换DOCX→PDF（{processed_count + 1}/{total}）：{f}"
                ))

                # 构建目标PDF路径
                if use_separate_folder:
                    relative_path = os.path.relpath(docx_path, root_dir)
                    pdf_path = os.path.join(root_dir, "docx2pdf", f"{os.path.splitext(relative_path)[0]}.pdf")
                    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
                else:
                    pdf_path = f"{os.path.splitext(docx_path)[0]}.pdf"

                success_flag = False
                error_msg = ""
                # 处理目标文件已存在的情况（单线程内询问，避免多线程弹窗冲突）
                if os.path.exists(pdf_path):
                    error_msg = "跳过：目标PDF已存在"
                else:
                    try:
                        doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
                        doc.ExportAsFixedFormat(
                            OutputFileName=os.path.abspath(pdf_path),
                            ExportFormat=17,  # 17=PDF格式
                            IncludeDocProps=True,
                            CreateBookmarks=1,  # 保留大纲书签
                            DocStructureTags=True
                        )
                        doc.Close(SaveChanges=0)
                        success_flag = True
                    except Exception as e:
                        error_msg = f"转换失败：{str(e)}"

                # 线程安全更新统计数据
                with progress_lock:
                    global processed_count, success_count, error_list
                    processed_count += 1
                    if success_flag:
                        success_count += 1
                    if error_msg:
                        error_list.append(f"{filename} - {error_msg}")
        finally:
            # 必须关闭Word进程，释放资源
            word.Quit()

    # 4. 启动线程池执行任务
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_threads) as executor:
        executor.map(worker, task_queues)

    # 5. 任务完成后显示结果
    root.after(0, lambda: show_convert_result("DOCX→PDF", total, use_separate_folder))


# ------------------------------
# 辅助功能触发函数（替换原函数）
# ------------------------------
def convert_doc_action():
    """触发DOC转DOCX功能（并行版）"""
    folder = folder_var.get().replace("已选择：", "")
    if not folder or folder == "等待选择文件夹...":
        messagebox.showwarning("警告", "请先选择文件夹")
        return
    # 禁用按钮，避免重复点击
    convert_doc_btn.config(state=tk.DISABLED)
    status_var.set("准备并行转换DOC→DOCX...")
    # 启动子线程执行并行转换（守护线程，避免程序退出残留）
    threading.Thread(
        target=parallel_convert_doc_to_docx,
        args=(folder, options['keep_source_doc'].get(), status_var),
        daemon=True
    ).start()


def convert_pdf_action():
    """触发DOCX转PDF功能（并行版）"""
    folder = folder_var.get().replace("已选择：", "")
    if not folder or folder == "等待选择文件夹...":
        messagebox.showwarning("警告", "请先选择文件夹")
        return
    # 禁用按钮，避免重复点击
    convert_pdf_btn.config(state=tk.DISABLED)
    status_var.set("准备并行转换DOCX→PDF...")
    # 启动子线程执行并行转换（守护线程，避免程序退出残留）
    threading.Thread(
        target=parallel_convert_docx_to_pdf,
        args=(folder, options['docx2pdf_separate_folder'].get(), status_var),
        daemon=True
    ).start()


# ------------------------------
# 主界面
# ------------------------------
def main():
    global options, root, process_btn, status_var, folder_var, convert_doc_btn, convert_pdf_btn
    root = tk.Tk()
    root.title("Word文件处理工具（全功能并行版）")
    root.geometry("700x800")
    root.resizable(False, False)

    # 变量定义
    folder_var = tk.StringVar(value="等待选择文件夹...")
    options = {
        # 主功能选项
        'remove_header_footer': tk.BooleanVar(value=True),
        'add_custom_header': tk.BooleanVar(value=True),
        'add_page_number': tk.BooleanVar(value=True),
        'replace_patterns': tk.BooleanVar(value=True),
        'set_question_outline': tk.BooleanVar(value=True),
        'keep_backup': tk.BooleanVar(value=False),
        # 辅助功能选项
        'keep_source_doc': tk.BooleanVar(value=False),
        'docx2pdf_separate_folder': tk.BooleanVar(value=False)
    }
    status_var = tk.StringVar(value="就绪")

    # 主框架
    main_frame = ttk.Frame(root, padding=15)
    main_frame.pack(fill=tk.BOTH, expand=True)

    # 文件夹选择区
    ttk.Label(main_frame, text="工作文件夹:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
    ttk.Label(main_frame, textvariable=folder_var, wraplength=650).pack(anchor=tk.W, pady=(0, 10))

    def select_folder_action():
        """选择文件夹并更新显示"""
        folder = select_folder()
        if folder:
            folder_var.set(f"已选择：{folder}")

    ttk.Button(main_frame, text="选择文件夹", command=select_folder_action).pack(anchor=tk.W, pady=(0, 20))

    # ------------------------------
    # 主功能区：Word处理
    # ------------------------------
    ttk.Separator(main_frame, orient="horizontal").pack(fill=tk.X, pady=10)
    ttk.Label(main_frame, text="【主功能区：Word文件并行处理】", font=("Arial", 11, "bold")).pack(anchor=tk.W,
                                                                                               pady=(0, 10))

    # 处理内容选项
    ttk.Label(main_frame, text="处理内容选项:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
    options_frame = ttk.Frame(main_frame)
    options_frame.pack(fill=tk.X, pady=(0, 10))

    # 选项列1
    col1 = ttk.Frame(options_frame)
    col1.pack(side=tk.LEFT, fill=tk.X, expand=True)
    ttk.Checkbutton(col1, text="删除页眉页脚", variable=options['remove_header_footer']).pack(anchor=tk.W, pady=2)
    ttk.Checkbutton(col1, text="添加自定义页眉", variable=options['add_custom_header']).pack(anchor=tk.W, pady=2)
    ttk.Checkbutton(
        col1,
        text="将题型、知识点、考点等段落大纲级别设置为1级",
        variable=options['set_question_outline']
    ).pack(anchor=tk.W, pady=2)

    # 选项列2
    col2 = ttk.Frame(options_frame)
    col2.pack(side=tk.LEFT, fill=tk.X, expand=True)
    ttk.Checkbutton(
        col2,
        text="添加居中页码（格式：第X页/共Y页）",
        variable=options['add_page_number']
    ).pack(anchor=tk.W, pady=2)
    ttk.Checkbutton(
        col2,
        text="替换指定文本模式（中英文括号/中括号）",
        variable=options['replace_patterns']
    ).pack(anchor=tk.W, pady=2)

    # 保存选项
    ttk.Label(main_frame, text="文件保存选项:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
    ttk.Checkbutton(
        main_frame,
        text="保留原文件为.bak备份（不勾选则直接替换源文件）",
        variable=options['keep_backup']
    ).pack(anchor=tk.W, pady=(0, 15))

    # 状态显示区
    ttk.Label(main_frame, textvariable=status_var, wraplength=650).pack(anchor=tk.W, pady=(0, 10))

    # 处理按钮（全局变量，用于禁用/启用）
    process_btn = ttk.Button(main_frame, text="开始并行处理Word文件", command=process_word_files_action)
    process_btn.pack(pady=(0, 15))

    # ------------------------------
    # 辅助功能区：格式转换
    # ------------------------------
    ttk.Separator(main_frame, orient="horizontal").pack(fill=tk.X, pady=10)
    ttk.Label(main_frame, text="【辅助功能区：格式转换（并行版）】", font=("Arial", 11, "bold")).pack(anchor=tk.W,
                                                                                                 pady=(0, 10))

    # DOC转DOCX
    ttk.Label(main_frame, text="DOC转DOCX选项:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
    ttk.Checkbutton(
        main_frame,
        text="保留源文件（.doc）",
        variable=options['keep_source_doc']
    ).pack(anchor=tk.W, pady=(0, 5))

    convert_doc_btn = ttk.Button(main_frame, text="并行批量转换DOC→DOCX", command=convert_doc_action)
    convert_doc_btn.pack(pady=(0, 10))

    # DOCX转PDF
    ttk.Label(main_frame, text="DOCX转PDF选项:", font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 5))
    ttk.Checkbutton(
        main_frame,
        text="保存到docx2pdf文件夹（不勾选则保存到原位置）",
        variable=options['docx2pdf_separate_folder']
    ).pack(anchor=tk.W, pady=(0, 5))

    convert_pdf_btn = ttk.Button(main_frame, text="并行批量转换DOCX→PDF", command=convert_pdf_action)
    convert_pdf_btn.pack(pady=(0, 10))

    # 退出按钮
    ttk.Button(main_frame, text="退出", command=lambda: [root.destroy(), os._exit(0)]).pack(pady=15)

    root.mainloop()


if __name__ == "__main__":
    main()